"""
Generate RFP for Ovivo Water at 2546 S Fry St, Boise ID
Site #6 — Two suites of 7,062 SF each = 14,124 SF total
"""
import copy
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

# ── deal parameters ──────────────────────────────────────────────────────────
PARAMS = {
    "DATE":               "March 11, 2026",
    "ADDRESS":            "2546 S Fry St, Boise, ID 83705",
    "TENANT_NAME":        "Ovivo Water",
    "LANDLORD_NAME":      "_____________________",
    "DEAR":               "_____________________",
    "RESPONSE_DEADLINE":  "March 18, 2026",
    "EXPIRY_DATE":        "March 25, 2026",

    # Premises
    "TOTAL_SF":           "14,124",
    "WH_SF":              "12,924",
    "OFFICE_SF":          "1,200",
    "PARKING_STALLS":     "17",

    # Dates / Term
    "COMMENCE_MONTHS":    "Zero (0)",
    "DELIVERY_DATE_TEXT": "immediately upon execution of the Lease. Landlord represents the Premises are currently available.",
    "LEASE_TERM":         "Sixty (60) calendar months",

    # Economics
    "BASE_RENT_PSF_ANNUAL": "$13.20",
    "BASE_RENT_BASIS":      "NNN",
    "ESCALATION_PCT":       "three percent (3%)",
    "OPEX_TAX":     "$0.10",
    "OPEX_INS":     "$0.04",
    "OPEX_CAM":     "$0.05",
    "OPEX_TOTAL":   "$0.19",

    # Renewal options
    "RENEWAL_OPTIONS": "two (2)",
}

# ── rent schedule ─────────────────────────────────────────────────────────────
sf   = 14124
rate = 1.10
esc  = 1.03
rent_schedule_lines = []
for yr in range(1, 6):
    monthly = rate * sf
    annual  = monthly * 12
    rent_schedule_lines.append(
        f"Year {yr}: ${rate:.4f}/SF/Mo  |  ${monthly:,.2f}/Mo  |  ${annual:,.2f}/Yr"
    )
    rate = round(rate * esc, 4)

RENT_SCHEDULE = "\n".join(rent_schedule_lines)


def replace_in_runs(para, old, new):
    """Replace 'old' with 'new' across all runs in a paragraph, preserving run formatting."""
    full_text = "".join(r.text for r in para.runs)
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    # Put new text in first run, clear the rest
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    return True


def replace_all_in_para(para, replacements):
    full_text = "".join(r.text for r in para.runs)
    changed = False
    for old, new in replacements.items():
        if old in full_text:
            full_text = full_text.replace(old, new)
            changed = True
    if changed and para.runs:
        para.runs[0].text = full_text
        for r in para.runs[1:]:
            r.text = ""
    return changed


def set_cell_text_preserve_fmt(cell, new_text):
    """Replace cell content with new_text, keeping the first paragraph's run formatting."""
    # Clear all paragraphs except the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    para = cell.paragraphs[0]
    # Clear runs
    for r in para.runs:
        r.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def add_para_to_cell(cell, text, bold=False):
    """Append a new paragraph to a cell."""
    p = cell.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True


# ── load template ─────────────────────────────────────────────────────────────
doc = Document("TEMPLATE - Industrial Long Form RFP.docx")

# ── fill header paragraphs ────────────────────────────────────────────────────
para_replacements = {
    "DATE":               PARAMS["DATE"],
    "ADDRESS":            PARAMS["ADDRESS"],
    "_________,":         f'{PARAMS["DEAR"]},',
    "___________________": PARAMS["TENANT_NAME"],
    # response deadline (first blank in that sentence)
    "__________________":  PARAMS["RESPONSE_DEADLINE"],
}

# Track which ____ we're on for the two separate date blanks
deadline_done = False
expiry_done   = False

for para in doc.paragraphs:
    t = "".join(r.text for r in para.runs)

    if t.strip() == "DATE":
        if para.runs:
            para.runs[0].text = PARAMS["DATE"]
            for r in para.runs[1:]:
                r.text = ""
        continue

    if "Request for Proposal at ADDRESS" in t:
        new = t.replace("ADDRESS", PARAMS["ADDRESS"])
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        continue

    if "_________," in t:
        new = t.replace("_________,", f'{PARAMS["DEAR"]},')
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        continue

    if '___________________' in t and '"Tenant"' in t:
        new = t.replace("___________________", PARAMS["TENANT_NAME"])
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        continue

    if "__________________" in t and "respond" in t and not deadline_done:
        new = t.replace("__________________", PARAMS["RESPONSE_DEADLINE"])
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        deadline_done = True
        continue

    if "__________________" in t and "appreciated" in t and not expiry_done:
        new = t.replace("__________________", PARAMS["EXPIRY_DATE"])
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        expiry_done = True
        continue

# ── fill table ────────────────────────────────────────────────────────────────
table = doc.tables[0]

def cell(r, c):
    return table.rows[r].cells[c]

def set_cell(r, c, text):
    cl = cell(r, c)
    # wipe existing paragraphs
    tc = cl._tc
    for p in tc.findall(qn('w:p')):
        tc.remove(p)
    # add fresh paragraph with run
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    tc.append(new_p)


# Row 0: Tenant
set_cell(0, 1, PARAMS["TENANT_NAME"])

# Row 1: Landlord
set_cell(1, 1, PARAMS["LANDLORD_NAME"])

# Row 2: Property Address
set_cell(2, 1, PARAMS["ADDRESS"])

# Row 3: Premises
set_cell(3, 1,
    f'{PARAMS["TOTAL_SF"]} rentable square feet comprised of two (2) contiguous suites '
    f'of approximately 7,062 SF each. Approximately {PARAMS["WH_SF"]} square feet of '
    f'warehouse/distribution space and approximately {PARAMS["OFFICE_SF"]} square feet '
    f'of office space.\n\n'
    f'Please include a floorplan of the Premises as Exhibit A, a building site plan with '
    f'your response to this proposal, and a separate AutoCAD file.'
)

# Row 4: Parking
set_cell(4, 1,
    f'Tenant shall also have access to approximately {PARAMS["PARKING_STALLS"]} dedicated '
    f'parking spaces of the exterior parking lot (pro-rata share of 102 total stalls based '
    f'on Tenant\'s proportionate occupancy of {PARAMS["TOTAL_SF"]} SF of the '
    f'84,750 SF building). Tenant may park vehicles in the parking stalls 24 hours per day, '
    f'365 days per year. Tenant shall have the right to designate certain stalls for '
    f'exclusive use by its employees and visitors.'
)

# Row 7: Commencement Date
set_cell(7, 1,
    f'{PARAMS["COMMENCE_MONTHS"]} Months following the Delivery Date, as defined below. '
    f'Tenant\'s Lease shall commence upon delivery of the Premises in the Delivery Condition.'
)

# Row 8: Delivery Date
set_cell(8, 1,
    f'The date upon which Landlord has delivered the Premises to Tenant in the Delivery '
    f'Condition described below. The Delivery Date shall be {PARAMS["DELIVERY_DATE_TEXT"]}'
)

# Row 10: Lease Term
set_cell(10, 1,
    f'{PARAMS["LEASE_TERM"]}, beginning on the Commencement Date.'
)

# Row 11: Base Rent
set_cell(11, 1,
    f'{PARAMS["BASE_RENT_PSF_ANNUAL"]} per square foot per year on a '
    f'{PARAMS["BASE_RENT_BASIS"]} basis ($1.10/SF/Mo in Year 1). Thereafter, Base Rent '
    f'shall be adjusted annually (including during each Renewal Term) with '
    f'{PARAMS["ESCALATION_PCT"]} annual increases commencing on each anniversary of the '
    f'Lease Commencement Date.\n\n'
    f'Projected Rent Schedule ({PARAMS["TOTAL_SF"]} RSF):\n'
    f'{RENT_SCHEDULE}'
)

# Row 12: Additional Rent
set_cell(12, 1,
    f'Please provide a detailed itemization of all Additional Rent for which Tenant will '
    f'be responsible to pay. Based on the tour book, Tenant anticipates OPEX of '
    f'approximately {PARAMS["OPEX_TOTAL"]}/SF/Mo. Please confirm the following breakdown '
    f'and provide 2024 actuals and 2025 estimates:\n\n'
    f'\u2022  Real Estate Taxes and Special Assessments: {PARAMS["OPEX_TAX"]}/SF/Mo (est.)\n'
    f'\u2022  Insurance: {PARAMS["OPEX_INS"]}/SF/Mo (est.)\n'
    f'\u2022  Exterior Area Maintenance (CAM): {PARAMS["OPEX_CAM"]}/SF/Mo (est.)\n\n'
    f'Please provide 2024 actual expenses and 2025 budget estimates with caps on '
    f'controllable expenses.'
)

# Row 17: Option to Renew
set_cell(17, 1,
    f'Tenant shall have {PARAMS["RENEWAL_OPTIONS"]} renewal options, each for a period of '
    f'five (5) years (each, a "Renewal Term") at the then-current Base Rent (as increased '
    f'annually per the escalation schedule above). Tenant shall provide at least six (6) '
    f'months\' advance written notice to Landlord of the exercise of any renewal option.'
)

# ── save output ───────────────────────────────────────────────────────────────
outfile = "RFP - Ovivo Water - 2546 S Fry St - Boise ID.docx"
doc.save(outfile)
print(f"Saved: {outfile}")
